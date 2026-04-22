import datetime
import io
import json
import os
import re
import traceback
from collections import defaultdict
from functools import wraps

import identity
import markdown
import webvtt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import (
    Blueprint,
    Response,
    jsonify,
    redirect,
    render_template,
    request,
    session,
    url_for,
    send_file,
)
from werkzeug.utils import secure_filename

from config import get_config
from custom_logger import logger
from utils.account_sql_utils import AccountSQLUtilities
from utils.auth_utils import require_login
# from utils.azure_ai_search_utils import AzureAiSearchUtilities
from utils.azure_utils import AzureBlobUtilities
from utils.chat_template_utilities import chat_template_utils
from utils.file_utils import fileUtilities
from utils.helper_utils import generate_uuid, get_greeting_message, get_model_details
from utils.llm_utils import LLMUtilities
from utils.session_group_sql_utils import session_group_sql_utils
from utils.session_items_utils import session_items_utils
from utils.sharepoint_utilities import sharepointUtilities
from flask import Flask, request, send_file
from markdown import markdown
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
from custom_agents.wrapper import UserInfo


from . import docbuild_bp

config = get_config()
llm_utils = LLMUtilities()
file_utils = fileUtilities()
azure_utils = AzureBlobUtilities()
account_utils = AccountSQLUtilities()
sharepoint_utils = sharepointUtilities()


@docbuild_bp.route("/", defaults={"session_id": None})
@docbuild_bp.route("/<session_id>")
@require_login
def chat_home(session_id):
    """Chat home page with optional session loading."""
    try:
        user_data = {
            "name": session.get("user"),
            "email": session.get("email"),
            "is_admin": account_utils.is_admin_user(session.get("email")),
            "profile_picture_path": session.get("profile_picture_path"),
        }

        greeting_message = get_greeting_message(session["user"])
        model_details = get_model_details()

        user_preferences = account_utils.get_user_by_user_id(session.get("user_id"))
        theme = "light" if not user_preferences or user_preferences.get('is_dark_mode_enabled') == 0 else "dark"

        return render_template(
            "docbuild_home.html",
            session_id=session_id,
            theme=theme,
            user=user_data,
            model_details=model_details,
            greeting_message=greeting_message,
            version=identity.__version__,
        )
    except Exception as e:
        logger.error(f"Error in chat_home: {str(e)}")
        return jsonify({"error": "An internal server error occurred"}), 500

@docbuild_bp.route("/update_session", methods=["POST"])
def update_session():
    """Updates session with new data."""
    try:
        data = request.get_json()
        session_id = data.get("session_id")
        prompt = data.get("prompt")
        chatbot_response = data.get("chatbot_response")
        session_item_message = data.get("session_item_message", "")
        new_messages_trace = data.get("new_messages_trace", "")
        session_item_id = data.get("session_item_id", "")

        logger.info("Received request to update session.")
        logger.info(data)
        logger.info(session_id)
        logger.info(prompt)
        logger.info(chatbot_response)

        user_id = session["user_id"]
        app_name = "docbuild"
        if chatbot_response != "":
            session_items_utils.insert_message(
                session_id=session_id,
                prompt=prompt,
                chatbot_response=chatbot_response,
                user_id=user_id,
                app_name=app_name,
                session_name=prompt,
                session_item_details={
                    "session_item_message": session_item_message,
                    "new_messages_trace": new_messages_trace,
                },
                session_item_id=session_item_id
            )

        return jsonify({"success": True})
    except Exception:
        logger.error("Error in update_session function:")
        logger.error(traceback.format_exc())
        return jsonify({"error": "An internal server error occurred"}), 500
    
@docbuild_bp.route("/get_session_data", methods=["GET"])
def get_session_data():
    """Retrieves session data."""
    try:
        session_id = request.args.get("session_id")

        session_data = session_items_utils.get_session_data(session_id, "docbuild")
        return jsonify(session_data)
    except Exception:
        logger.error("Error in get_session_data function:")
        logger.error(traceback.format_exc())
        return jsonify({"error": "An internal server error occurred"}), 500

@docbuild_bp.route('/download_markdown_as_word', methods=['POST'])
def download_markdown_as_word():
    """
    Handle requests to convert Markdown to a Word document and send it as a downloadable file.
    """
    try:
        # Extract Markdown content from the request JSON
        markdown_content = request.json.get("markdown", "")

        # Convert Markdown to HTML
        html_content = markdown(markdown_content)

        # Parse the HTML using BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        # Create a Word document
        doc = Document()

        # Process each HTML tag and add elements to the Word document
        for tag in soup.descendants:
            if tag.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(tag.name[1])  # Extract heading level (e.g., 1 for <h1>)
                doc.add_heading(tag.text, level=level)
            elif tag.name == "p":
                doc.add_paragraph(tag.text)
            elif tag.name in ["ul", "ol"]:
                for li in tag.find_all("li"):
                    doc.add_paragraph(f"• {li.text}", style="List Bullet")

        # Save the Word document to an in-memory buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Return the Word file as a downloadable response
        return send_file(
            buffer,
            as_attachment=True,
            download_name="document.docx",  # Filename for the downloaded file
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        return {"error": str(e)}, 500
@docbuild_bp.route('/download_html_as_file', methods=['POST'])
def download_html_as_file():
    try:
        html_content = request.json.get("markdown", "")
        buffer = BytesIO()
        buffer.write(html_content.encode("utf-8"))
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="document.html",
            mimetype="text/html",
        )
    except Exception as e:
        return {"error": str(e)}, 500
@docbuild_bp.route('/download_text_as_file', methods=['POST'])
def download_text_as_file():
    try:
        text_content = request.json.get("markdown", "")
        buffer = BytesIO()
        buffer.write(text_content.encode("utf-8"))
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="document.txt",
            mimetype="text/plain",
        )
    except Exception as e:
        return {"error": str(e)}, 500
@docbuild_bp.route('/download_json_as_file', methods=['POST'])
def download_json_as_file():
    try:
        json_text = request.json.get("markdown", "")
        buffer = BytesIO()
        buffer.write(json_text.encode("utf-8"))
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="document.json",
            mimetype="application/json",
        )
    except Exception as e:
        return {"error": str(e)}, 500
@docbuild_bp.route('/download_code_as_file', methods=['POST'])
def download_code_as_file():
    try:
        code_text = request.json.get("markdown", "")
        filetype = request.args.get("ext", "txt")  # you can pass ?ext=py, js, ts etc.
        buffer = BytesIO()
        buffer.write(code_text.encode("utf-8"))
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"script.{filetype}",
            mimetype="text/plain",
        )
    except Exception as e:
        return {"error": str(e)}, 500
@docbuild_bp.route('/download_csv_as_file', methods=['POST'])
def download_csv_as_file():
    try:
        csv_text = request.json.get("markdown", "")
        buffer = BytesIO()
        buffer.write(csv_text.encode("utf-8"))
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="document.csv",
            mimetype="text/csv",
        )
    except Exception as e:
        return {"error": str(e)}, 500
@docbuild_bp.route('/download_yaml_as_file', methods=['POST'])
def download_yaml_as_file():
    try:
        yaml_text = request.json.get("markdown", "")
        buffer = BytesIO()
        buffer.write(yaml_text.encode("utf-8"))
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="document.yaml",
            mimetype="text/yaml",
        )
    except Exception as e:
        return {"error": str(e)}, 500

@docbuild_bp.route("/generate", methods=["POST"])
def generate_document():
    try:
        data = request.get_json()
        user_input = data.get("prompt")
        session_id = data.get("session_id")
        model_name = data.get("model_name")
        current_document_text = data.get("current_document_text")

        user_id = session.get("user_id")

        # Log the received parameters
        logger.info("Received generation request:")
        logger.info(f"user_input: {user_input}")
        logger.info(f"Session ID: {session_id}")
        logger.info(f"model_name: {model_name}")
        logger.info(f"current_document_text: {current_document_text}")

        from apps.docbuild.src.prompt_templates import messages as default_messages
        
        previous_messages = session_items_utils.get_previous_messages(
            session_id, 
            "etexchat" # docbuild is also using the same logic as etexchat. Hence using etexchat.
        )

        messages = default_messages + previous_messages

        prompt = user_input.strip()
        if current_document_text.strip() != "":
            prompt = f"""{user_input}\n\nNote: Apply changes to the current file and give the whole file back\n\nCurrent file:\n\n{current_document_text}"""
        
        
        msg = {"role": "user", "content": [{"type": "input_text", "text": prompt}]}
        user_info = UserInfo(
            email=session.get("user_id"),
            auth_access_token="fake_token",
            user_id=user_id,
            session_id=session_id,
            entra_id_user_id= session['entra_id_user_id']
        )


        return Response(
            _stream_llm_response(messages + [msg], model_name, "Document Builder", user_info), 
            content_type="text/event-stream"
        )
            
    except Exception as e:
        logger.error(f"Error getting chatbot response: {e}")
        return jsonify({"message": "Error getting chatbot response"}), 500

def _stream_llm_response(messages, model_name, user_selected_agent, user_info):
    """
    Bridge function to handle async streaming in a synchronous Flask context.
    Uses a queue-based approach to stream data from async function.
    
    If any error occurs during streaming, the function raises the error and logs the traceback.
    It does not yield error messages to the client — it crashes cleanly for Flask to catch and handle.
    """
    from queue import Queue
    import threading
    import asyncio
    import traceback

    response_queue = Queue()
    exception_container = [None]

    def run_async_stream():
        """Run the async streaming function in a separate thread."""
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

            async def stream_data():
                try:
                    async for chunk in llm_utils.invoke_llm_stream(messages, model_name, user_selected_agent, user_info):
                        response_queue.put(('data', chunk))
                except Exception as e:
                    exception_container[0] = e
                    response_queue.put(('error', None))
                finally:
                    response_queue.put(('done', None))

            loop.run_until_complete(stream_data())

        except Exception as e:
            exception_container[0] = e
            response_queue.put(('error', None))
        finally:
            response_queue.put(('done', None))

    thread = threading.Thread(target=run_async_stream)
    thread.daemon = True
    thread.start()

    # Yield data or raise error immediately
    while True:
        try:
            event_type, data = response_queue.get(timeout=300)

            if event_type == 'data':
                yield data
            elif event_type == 'error':
                if exception_container[0]:
                    logger.error("Exception during streaming:")
                    logger.error(traceback.format_exc())
                    raise exception_container[0]
            elif event_type == 'done':
                break

        except Exception as e:
            logger.error("Unexpected error in streaming loop:")
            logger.error(traceback.format_exc())
            raise  # Let Flask catch and respond with 500